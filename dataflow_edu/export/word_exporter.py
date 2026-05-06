# -*- coding: utf-8 -*-
"""Word 导出（python-docx）。

排版规则（与 plan.md M3 对齐）：
- 章节层级：category（H1）→ subcategory（H2）→ type（H3），全局连续编号。
- variant='with_answer'：每题正文 + 答案 + 解析 inline 排版。
- variant='blank'：主体只有题目+选项；分页符之后单独一个「参考答案与解析」区域，
  按相同的章节层级再次列出每题答案+解析（编号沿用主体）。
- lang='zh'|'en'|'fr'：题目/答案/解析字段已在 data_loader 里 resolve 完毕，
  本模块只关心排版。

题号规则：
- 同一 subcategory 内连续编号（1, 2, 3…），换新 subcategory 时重置为 1。
- 每个题号带题型缩写前缀，如「单选-1」「简答-3」；无法识别的题型取前 2 字。

目录规则：
- 在封面标题之后插入真·Word TOC 字段（1-3 级标题），dirty=true 以便 Word/LibreOffice
  打开或转换时刷新页码和超链接。
- 字段结果里同时写入一份静态兜底目录，避免无字段刷新能力的 PDF 转换环境输出空目录。
- 目录之后加分页符，确保正文从新页开始。

中文字体在 docx 里非常坑：默认 Calibri 对中文不友好。这里强制全文（中文 east-asia
+ 西文 ascii/hAnsi）统一为「微软雅黑」，并同步覆盖 Title / Heading 1~3 等内置样式，
保证 docx 与下游 LibreOffice 转出的 PDF 字体一致。
"""

from __future__ import annotations

import os
import shutil
import socket
import subprocess
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from dataflow_edu.export.data_loader import (
    QuestionRecord,
    load_task_questions,
)


VARIANT_WITH_ANSWER = "with_answer"
VARIANT_BLANK = "blank"
SUPPORTED_VARIANTS = (VARIANT_WITH_ANSWER, VARIANT_BLANK)

# UI 上显示的章节标题文案（按语言）
_HEADINGS = {
    "zh": {
        "title": "试题",
        "answers_section": "参考答案与解析",
        "uncategorized": "未分类",
        "answer_label": "【答案】",
        "explanation_label": "【解析】",
        "type_label": "题型",
        "difficulty_label": "难度",
        "page_label": "出处",
        "toc_title": "目录",
    },
    "en": {
        "title": "Questions",
        "answers_section": "Answers & Explanations",
        "uncategorized": "Uncategorized",
        "answer_label": "[Answer] ",
        "explanation_label": "[Explanation] ",
        "type_label": "Type",
        "difficulty_label": "Difficulty",
        "page_label": "Source",
        "toc_title": "Table of Contents",
    },
    "fr": {
        "title": "Questions",
        "answers_section": "Réponses et explications",
        "uncategorized": "Non classé",
        "answer_label": "[Réponse] ",
        "explanation_label": "[Explication] ",
        "type_label": "Type",
        "difficulty_label": "Difficulté",
        "page_label": "Source",
        "toc_title": "Table des matières",
    },
}


DEFAULT_FONT = os.environ.get("DATAFLOW_EDU_EXPORT_FONT", "").strip() or "微软雅黑"

# 需要统一字体的内置样式（Title 是封面大标题；Heading 1~3 是章节；TOC 是目录）
_STYLES_TO_UNIFY = (
    "Normal",
    "Title",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "TOC Heading",
    "TOC 1",
    "TOC 2",
    "TOC 3",
)

# 常见题型 → 缩写映射；未命中时截取前 2 个字
_TYPE_ABBREV: Dict[str, str] = {
    "单项选择题": "单选",
    "多项选择题": "多选",
    "单选题": "单选",
    "多选题": "多选",
    "判断题": "判断",
    "填空题": "填空",
    "简答题": "简答",
    "论述题": "论述",
    "计算题": "计算",
    "实验题": "实验",
    "材料分析题": "材料",
    "综合题": "综合",
    "案例分析题": "案例",
    "分析题": "分析",
    "应用题": "应用",
    "证明题": "证明",
}

QuestionItems = List[Tuple[str, QuestionRecord]]
TypeGroup = List[Tuple[str, QuestionItems]]
SubcategoryGroup = List[Tuple[str, TypeGroup]]
GroupedQuestions = List[Tuple[str, SubcategoryGroup]]


def _abbrev_type(typ: str) -> str:
    """返回题型简写：优先查表，否则截取前 2 个字符。"""
    if not typ or typ == "__uncat__":
        return ""
    return _TYPE_ABBREV.get(typ, typ[:2] if len(typ) > 2 else typ)


def _make_label(typ: str, sub_local_idx: int) -> str:
    """生成带题型前缀的题号，如 '单选-1'、'简答-3'。无可识别前缀时直接返回数字字符串。"""
    prefix = _abbrev_type(typ)
    if prefix:
        return f"{prefix}-{sub_local_idx}"
    return str(sub_local_idx)


def _set_rfonts(rPr: Any, font_name: str) -> None:
    """Clear theme font refs and pin all OOXML font slots to font_name."""

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def _set_run_font(run: Any, font_name: str) -> None:
    """Set a concrete font on a run, including the eastAsia slot."""

    run.font.name = font_name
    _set_rfonts(run._r.get_or_add_rPr(), font_name)


def _force_style_font(doc: DocxDocument, style_name: str, font_name: str) -> None:
    """把指定样式的中英文字体都改成 font_name。

    重点坑：python-docx 默认模板里 Title / Heading 1~3 的 rFonts 用的是
    **主题字体引用**（`asciiTheme="majorHAnsi"` / `eastAsiaTheme="majorEastAsia"` 等）。
    OOXML 里 `*Theme` 优先级高于具体名，不清掉它们，写多少具体字体名都没用——
    渲染器会去 theme1.xml 查主题字体，默认主题里 `majorEastAsia` 没设中文字体，
    Windows 上就会回落到 MS Gothic / MS Mincho。
    所以这里：1) 先删 4 个 *Theme 属性；2) 再写 4 个具体 slot。
    """

    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    rPr = style.element.get_or_add_rPr()
    _set_rfonts(rPr, font_name)
    # python-docx 的 style.font.name 只会写 ascii / hAnsi，这里补一刀方便后续读取
    style.font.name = font_name


def _force_default_font(doc: DocxDocument, font_name: str) -> None:
    """Set w:docDefaults so LibreOffice does not fall back to theme fonts."""

    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_elem.insert(0, doc_defaults)

    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)

    rPr = rpr_default.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rpr_default.append(rPr)
    _set_rfonts(rPr, font_name)


def _set_east_asia_font(doc: DocxDocument, font_name: str = DEFAULT_FONT) -> None:
    """把所有相关样式的中英文字体统一为 font_name（默认微软雅黑）。"""

    _force_default_font(doc, font_name)
    for style_name in _STYLES_TO_UNIFY:
        _force_style_font(doc, style_name, font_name)
    doc.styles["Normal"].font.size = Pt(11)


def _grouped(
    records: List[QuestionRecord],
) -> GroupedQuestions:
    """按 category → subcategory → type 三级分组，最内层附上显示标签。

    标签格式为「题型缩写-subcategory内序号」，如 '单选-1'。
    同一 subcategory 内题号连续递增（跨 type），换 subcategory 时重置为 1。

    返回结构：
        [(category, [(subcategory, [(type, [(label, record), ...]), ...]), ...]), ...]
    """

    raw: "Dict[str, Dict[str, Dict[str, List[QuestionRecord]]]]" = {}
    for r in records:
        cat = r.category or "__uncat__"
        sub = r.subcategory or "__uncat__"
        typ = r.type or "__uncat__"
        raw.setdefault(cat, {}).setdefault(sub, {}).setdefault(typ, []).append(r)

    result = []
    for cat, sub_map in raw.items():
        sub_list = []
        for sub, type_map in sub_map.items():
            type_list: TypeGroup = []
            sub_local = 1
            for typ, recs in type_map.items():
                items: QuestionItems = []
                for rec in recs:
                    items.append((_make_label(typ, sub_local), rec))
                    sub_local += 1
                type_list.append((typ, items))
            sub_list.append((sub, type_list))
        result.append((cat, sub_list))
    return result


def _normalize_label(label: str, fallback: str) -> str:
    return fallback if label == "__uncat__" else label


def _format_options(options: Any) -> List[str]:
    """返回每个选项的渲染字符串。兼容 list[str] / dict[letter -> text]。"""

    if not options:
        return []
    if isinstance(options, list):
        rendered = []
        existing_prefixes = {
            f"{chr(c)}{sep}" for c in range(ord("A"), ord("H")) for sep in (".", "、", ":", "：")
        }
        for i, opt in enumerate(options):
            text = "" if opt is None else str(opt)
            stripped = text.lstrip()
            if stripped[:2] in existing_prefixes:
                rendered.append(stripped)
            else:
                rendered.append(f"{chr(ord('A') + i)}. {stripped}")
        return rendered
    if isinstance(options, dict):
        rendered = []
        for letter in sorted(options.keys()):
            text = options[letter]
            rendered.append(f"{letter}. {text}")
        return rendered
    return [str(options)]


def _add_question_body(doc: DocxDocument, label: str, rec: QuestionRecord) -> None:
    """渲染单题的题干 + 选项（不含答案/解析）。label 如 '单选-1'。"""

    p = doc.add_paragraph()
    p.add_run(f"{label}. ").bold = True
    p.add_run(rec.question or "")

    for line in _format_options(rec.options):
        op = doc.add_paragraph(line)
        op.paragraph_format.left_indent = Pt(18)


def _add_question_answer(
    doc: DocxDocument,
    label: str,
    rec: QuestionRecord,
    *,
    labels: Dict[str, str],
    include_question_repeat: bool,
) -> None:
    """渲染单题的答案+解析。

    `include_question_repeat=True` 时（学生卷答案区），先重复一行题号 + 题干截断，
    便于学生对照；with_answer 卷正文已经有题干，无需重复。
    """

    if include_question_repeat:
        head = doc.add_paragraph()
        head.add_run(f"{label}. ").bold = True
        snippet = (rec.question or "").strip().replace("\n", " ")
        if len(snippet) > 60:
            snippet = snippet[:60] + "…"
        head.add_run(snippet)

    if rec.answer:
        p = doc.add_paragraph()
        p.add_run(labels["answer_label"]).bold = True
        p.add_run(rec.answer)
    if rec.explanation:
        p = doc.add_paragraph()
        p.add_run(labels["explanation_label"]).bold = True
        p.add_run(rec.explanation)


def _enable_update_fields_on_open(doc: DocxDocument) -> None:
    """在文档 settings.xml 中写入 updateFields=true。

    Word 打开含此设置的文档时会自动刷新所有字段（含 TOC），
    无需用户手动 Ctrl+A → F9。
    """
    settings_elem = doc.settings.element
    # 避免重复写入
    existing = settings_elem.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_elem.append(update_fields)


def _toc_count_suffix(count: int, lang: str) -> str:
    if lang == "zh":
        return f"（{count}题）"
    if lang == "fr":
        return f" ({count} questions)"
    return f" ({count} questions)"


def _count_type_items(type_list: TypeGroup) -> int:
    return sum(len(items) for _, items in type_list)


def _count_sub_items(sub_list: SubcategoryGroup) -> int:
    return sum(_count_type_items(type_list) for _, type_list in sub_list)


def _append_field_char(run: Any, field_type: str, *, dirty: bool = False) -> None:
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), field_type)
    if dirty:
        fld_char.set(qn("w:dirty"), "true")
    run._r.append(fld_char)


def _append_toc_instruction(run: Any) -> None:
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = r" TOC \f \h \z "
    run._r.append(instr)


def _append_tc_field(paragraph: Any, text: str, level: int) -> None:
    """在标题段落里追加隐藏 TC 字段，供 TOC 刷新后保留题量等自定义目录文案。"""

    safe_text = text.replace('"', "'")
    r_begin = paragraph.add_run()
    _append_field_char(r_begin, "begin")

    r_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f' TC "{safe_text}" \\l {level} '
    r_instr._r.append(instr)

    r_end = paragraph.add_run()
    _append_field_char(r_end, "end")


def _add_paragraph_if_style_exists(doc: DocxDocument, style_name: str) -> Any:
    try:
        doc.styles[style_name]
    except KeyError:
        return doc.add_paragraph()
    return doc.add_paragraph(style=style_name)


def _add_toc_result_paragraph(
    doc: DocxDocument,
    text: str,
    *,
    level: int,
    start_field: bool = False,
    end_field: bool = False,
) -> None:
    """写入 TOC 字段的静态兜底结果段落。

    正常情况下 Word/LibreOffice 会用真实 TOC 替换这些段落；如果转换环境不刷新字段，
    PDF 至少仍有完整层级和题量信息，不会只剩一个空字段或提示文本。
    """

    p = _add_paragraph_if_style_exists(doc, f"TOC {level}")

    p.paragraph_format.left_indent = Pt(18 * (level - 1))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if start_field:
        r_begin = p.add_run()
        _append_field_char(r_begin, "begin", dirty=True)
        r_instr = p.add_run()
        _append_toc_instruction(r_instr)
        r_sep = p.add_run()
        _append_field_char(r_sep, "separate")

    run = p.add_run(text)
    if level == 1:
        run.bold = True

    if end_field:
        r_end = p.add_run()
        _append_field_char(r_end, "end")


def _build_toc_fallback_entries(
    grouped: GroupedQuestions,
    *,
    lang: str,
    variant: str,
    labels: Dict[str, str],
) -> List[Tuple[int, str]]:
    entries: List[Tuple[int, str]] = []

    for cat, sub_list in grouped:
        cat_text = _normalize_label(cat, labels["uncategorized"])
        entries.append((1, f"{cat_text} {_toc_count_suffix(_count_sub_items(sub_list), lang)}"))
        for sub, type_list in sub_list:
            sub_text = _normalize_label(sub, labels["uncategorized"])
            entries.append(
                (2, f"{sub_text} {_toc_count_suffix(_count_type_items(type_list), lang)}")
            )
            for typ, items in type_list:
                typ_text = _normalize_label(typ, labels["uncategorized"])
                entries.append((3, f"{typ_text} {_toc_count_suffix(len(items), lang)}"))

    if variant == VARIANT_BLANK and grouped:
        entries.append((1, labels["answers_section"]))
        for cat, sub_list in grouped:
            cat_text = _normalize_label(cat, labels["uncategorized"])
            entries.append((2, cat_text))
            for sub, _ in sub_list:
                sub_text = _normalize_label(sub, labels["uncategorized"])
                entries.append((3, sub_text))

    return entries


def _insert_toc(
    doc: DocxDocument,
    grouped: GroupedQuestions,
    lang: str = "zh",
    variant: str = VARIANT_WITH_ANSWER,
    labels: Optional[Dict[str, str]] = None,
) -> None:
    """插入可更新的 Word TOC 字段，随后分页。"""

    if labels is None:
        labels = _HEADINGS.get(lang, _HEADINGS["zh"])
    toc_title = labels["toc_title"]

    # 目录标题
    p_title = _add_paragraph_if_style_exists(doc, "TOC Heading")
    p_title.paragraph_format.space_after = Pt(6)
    run = p_title.add_run(toc_title)
    run.bold = True
    run.font.size = Pt(14)

    entries = _build_toc_fallback_entries(
        grouped,
        lang=lang,
        variant=variant,
        labels=labels,
    )
    if not entries:
        entries = [(1, labels["uncategorized"])]

    for idx, (level, text) in enumerate(entries):
        _add_toc_result_paragraph(
            doc,
            text,
            level=level,
            start_field=idx == 0,
            end_field=idx == len(entries) - 1,
        )

    # 目录后分页
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _resolve_soffice_bin() -> Optional[str]:
    env = os.environ.get("LIBREOFFICE_BIN", "").strip()
    if env and Path(env).is_file():
        return env
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    if os.name == "nt":
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(candidate).is_file():
                return candidate
    return None


def _find_free_local_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return int(s.getsockname()[1])


def _refresh_toc_with_libreoffice(docx_path: Path, *, timeout: int = 90) -> bool:
    """用 LibreOffice 排版引擎刷新 TOC 页码并保存 docx。

    python-docx 不能计算页码；这里通过 UNO 打开文档、更新目录索引、保存。
    LibreOffice/UNO 不可用时返回 False，保留可打开时自动更新的字段。
    """

    soffice = _resolve_soffice_bin()
    if not soffice:
        return False
    try:
        import uno
        from com.sun.star.beans import PropertyValue
    except Exception:
        return False

    def prop(name: str, value: Any) -> Any:
        p = PropertyValue()
        p.Name = name
        p.Value = value
        return p

    with tempfile.TemporaryDirectory(prefix="edu_docx_refresh_") as tmp:
        profile_dir = (Path(tmp) / ".lo_profile").resolve()
        profile_dir.mkdir(parents=True, exist_ok=True)
        try:
            port = _find_free_local_port()
        except OSError:
            return False
        accept = f"socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext"
        cmd = [
            soffice,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--norestore",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            f"--accept={accept}",
        ]
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        try:
            local_ctx = uno.getComponentContext()
            resolver = local_ctx.ServiceManager.createInstanceWithContext(
                "com.sun.star.bridge.UnoUrlResolver",
                local_ctx,
            )
            ctx = None
            deadline = time.time() + timeout
            last_error: Optional[Exception] = None
            while time.time() < deadline:
                if proc.poll() is not None:
                    break
                try:
                    ctx = resolver.resolve(f"uno:{accept}")
                    break
                except Exception as exc:
                    last_error = exc
                    time.sleep(0.25)
            if ctx is None:
                if last_error:
                    return False
                return False

            desktop = ctx.ServiceManager.createInstanceWithContext(
                "com.sun.star.frame.Desktop",
                ctx,
            )
            doc = desktop.loadComponentFromURL(
                docx_path.resolve().as_uri(),
                "_blank",
                0,
                (prop("Hidden", True), prop("ReadOnly", False)),
            )
            if doc is None:
                return False
            try:
                indexes = doc.getDocumentIndexes()
                for i in range(indexes.getCount()):
                    indexes.getByIndex(i).update()
                try:
                    doc.getTextFields().refresh()
                except Exception:
                    pass
                doc.store()
            finally:
                doc.close(True)
            return True
        except Exception:
            return False
        finally:
            proc.terminate()
            try:
                proc.communicate(timeout=10)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.communicate()


_BRAND_FOOTER = (
    "北京中关村学院 & 中关村人工智能研究院  ·  DataFlow-EDU  ·  Empowering Education with AI"
)


def _add_brand_footer(doc: DocxDocument, font_name: str = DEFAULT_FONT) -> None:
    """在文档所有节的页脚中央插入品牌信息。"""

    for section in doc.sections:
        section.different_first_page_header_footer = False
        footer = section.footer
        # 清空默认段落内容
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""
        # 取第一段（docx 页脚至少有一个默认段落）
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(_BRAND_FOOTER)
        _set_run_font(run, font_name)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9E, 0xA3, 0xAF)  # slate-400


def export_word(
    task_dir: str | os.PathLike[str],
    output_path: str | os.PathLike[str],
    *,
    stage: str = "3_8_mcq_verified",
    lang: str = "zh",
    variant: str = VARIANT_WITH_ANSWER,
    task_name: Optional[str] = None,
) -> Path:
    """生成 .docx。"""

    if variant not in SUPPORTED_VARIANTS:
        raise ValueError(f"unsupported variant: {variant!r}, expected one of {SUPPORTED_VARIANTS}")
    labels = _HEADINGS.get(lang, _HEADINGS["zh"])

    records = load_task_questions(task_dir, stage=stage, lang=lang)
    grouped = _grouped(records)

    doc = Document()
    _set_east_asia_font(doc)

    title = task_name or labels["title"]
    if variant == VARIANT_BLANK:
        title = f"{title} ({labels['title']})"
    h = doc.add_heading(title, level=0)
    h.alignment = 1  # center

    # 目录（封面标题之后，正文之前）
    _insert_toc(doc, grouped, lang, variant, labels)

    # 题干主体
    for cat_idx, (cat, sub_list) in enumerate(grouped):
        # 第二个及之后的 category 前插分页符
        if cat_idx > 0:
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

        cat_text = _normalize_label(cat, labels["uncategorized"])
        cat_heading = doc.add_heading(cat_text, level=1)
        _append_tc_field(
            cat_heading,
            f"{cat_text} {_toc_count_suffix(_count_sub_items(sub_list), lang)}",
            1,
        )
        for sub, type_list in sub_list:
            sub_text = _normalize_label(sub, labels["uncategorized"])
            sub_heading = doc.add_heading(sub_text, level=2)
            _append_tc_field(
                sub_heading,
                f"{sub_text} {_toc_count_suffix(_count_type_items(type_list), lang)}",
                2,
            )
            for typ, items in type_list:
                typ_text = _normalize_label(typ, labels["uncategorized"])
                typ_heading = doc.add_heading(typ_text, level=3)
                _append_tc_field(
                    typ_heading,
                    f"{typ_text} {_toc_count_suffix(len(items), lang)}",
                    3,
                )
                for label, rec in items:
                    _add_question_body(doc, label, rec)
                    if variant == VARIANT_WITH_ANSWER:
                        _add_question_answer(
                            doc, label, rec, labels=labels, include_question_repeat=False
                        )
                    # 题与题间留个空行
                    doc.add_paragraph("")

    # 学生卷：分页 + 答案区
    if variant == VARIANT_BLANK and records:
        # 分页符（当前已有，保持）
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        # TC level=1 使「参考答案与解析」在目录中与 category 同级。
        answer_heading = doc.add_heading(labels["answers_section"], level=1)
        _append_tc_field(answer_heading, labels["answers_section"], 1)
        for cat, sub_list in grouped:
            cat_text = _normalize_label(cat, labels["uncategorized"])
            cat_heading = doc.add_heading(cat_text, level=2)
            _append_tc_field(cat_heading, cat_text, 2)
            for sub, type_list in sub_list:
                sub_text = _normalize_label(sub, labels["uncategorized"])
                sub_heading = doc.add_heading(sub_text, level=3)
                _append_tc_field(sub_heading, sub_text, 3)
                for typ, items in type_list:
                    for label, rec in items:
                        _add_question_answer(
                            doc, label, rec, labels=labels, include_question_repeat=True
                        )
                        doc.add_paragraph("")

    _add_brand_footer(doc)
    _enable_update_fields_on_open(doc)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    _refresh_toc_with_libreoffice(out)
    return out
